"""
Emmanuel's Documentation Agent
==============================
this runs AFTER the research agent finishes. it takes the raw JSON
results and turns them into a proper formatted document — like a
real research report you could actually hand to someone.

it uses gemini pro to write a proper human-readable summary from the
raw data points, then builds a Word doc (.docx) and a PDF with nice
formatting, tables, headers, the works.

the idea is:
  1. research agent browses the web and collects raw data
  2. documentation agent takes that data and makes it look professional
  3. you get a report you can actually send to your manager

usage:
    from agents.documentation_agent import DocumentationAgent
    doc_agent = DocumentationAgent(client)
    doc_agent.generate(research_result)           # uses the dict directly
    doc_agent.generate_from_file("result.json")   # or from a saved JSON

last updated: 28 feb 2026 - emmanuel
"""
import json
import subprocess
import sys
from datetime import datetime
from pathlib import Path
from typing import Optional

# try importing these, install if missing
try:
    from fpdf import FPDF
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "fpdf2", "--quiet"])
    from fpdf import FPDF

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "--quiet"])
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT


from core.settings import DEFAULT_MODEL

OUTPUT_DIR = Path("outputs/research_results")


class DocumentationAgent:
    """
    takes research agent output and makes it into a proper document.
    can optionally use gemini to rewrite the summary in better english.
    """

    def __init__(self, client=None):
        # client is optional — if you pass a gemini client it'll use
        # pro to write a better summary. if not it just uses what the
        # research agent already wrote.
        self.client = client

    def generate(self, result: dict, output_name: str = None) -> dict:
        """
        main method — give it a research result dict and it makes:
          - a Word doc (.docx)
          - a PDF report
        returns paths to both files.
        """
        query = result.get("query", "Unknown Query")

        # if we have a gemini client, ask pro to write a better summary
        # from the raw data points. makes the report way more readable.
        if self.client and result.get("findings", {}).get("data_points"):
            enhanced_summary = self._enhance_summary(result)
            if enhanced_summary:
                result["findings"]["enhanced_summary"] = enhanced_summary

        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

        # figure out filenames
        if output_name is None:
            safe = "".join(c if c.isalnum() or c == " " else "_" for c in query[:40])
            output_name = f"report_{safe.strip().replace(' ', '_')}_{datetime.now():%Y%m%d_%H%M%S}"

        # generate both formats
        docx_path = self._make_docx(result, OUTPUT_DIR / f"{output_name}.docx")
        pdf_path = self._make_pdf(result, OUTPUT_DIR / f"{output_name}.pdf")

        print(f"\n{'='*60}")
        print(f"DOCUMENTATION AGENT — done")
        print(f"  Word: {docx_path}")
        print(f"  PDF:  {pdf_path}")
        print(f"{'='*60}\n")

        return {"docx": str(docx_path), "pdf": str(pdf_path)}

    def generate_from_file(self, json_path: str) -> dict:
        """load a research JSON and generate docs from it."""
        result = json.loads(Path(json_path).read_text())
        return self.generate(result)

    def generate_from_latest(self) -> dict:
        """find the most recent research JSON and generate docs."""
        jsons = sorted(OUTPUT_DIR.glob("research_*.json"), key=lambda p: p.stat().st_mtime)
        if not jsons:
            print("no research results found — run the research agent first")
            return None
        print(f"Using: {jsons[-1].name}")
        result = json.loads(jsons[-1].read_text())
        return self.generate(result)

    # ─── SUMMARY ENHANCEMENT ─────────────────────────────────────────
    # this is the cool part — we use gemini pro (text only, cheap) to
    # rewrite the raw findings into a proper paragraph that actually
    # reads like a human wrote it.

    def _enhance_summary(self, result: dict) -> Optional[str]:
        """ask gemini pro to write a proper summary from the raw data."""
        try:
            findings = result["findings"]
            data_text = ""
            for dp in findings.get("data_points", []):
                data_text += f"- {dp.get('fact','')}: {dp.get('value','')} {dp.get('unit','')}\n"

            prompt = f"""You are a technical documentation writer. Based on these research findings, 
write a clear 2-3 paragraph summary suitable for a professional report.

RESEARCH QUERY: {result['query']}

RAW DATA POINTS:
{data_text}

ORIGINAL SUMMARY: {findings.get('summary', 'N/A')}

SOURCES: {', '.join(findings.get('sources', [])[:5])}

CONFIDENCE: {findings.get('confidence', 'unknown')}

Write the summary in a professional but readable tone. Include specific numbers 
and values. Mention the sources briefly. If confidence is low, note any gaps.
Do NOT use markdown formatting. Just plain text paragraphs."""

            resp = self.client.models.generate_content(
                model=DEFAULT_MODEL,
                contents=prompt
            )
            print("[DocumentationAgent] Enhanced summary generated")
            return resp.text
        except Exception as e:
            # if pro fails we just use the original summary, no big deal
            print(f"[DocumentationAgent] Couldnt enhance summary ({e}), using original")
            return None

    # ─── WORD DOCUMENT ───────────────────────────────────────────────
    # proper .docx with headers, tables, formatting — looks professional

    def _make_docx(self, result: dict, filepath: Path) -> Path:
        """build a proper Word document from the research results."""
        f = result["findings"]
        m = result.get("metadata", {})
        query = result.get("query", "Unknown")

        doc = Document()

        # --- page setup ---
        section = doc.sections[0]
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        # --- title ---
        title = doc.add_heading("Research Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in title.runs:
            run.font.color.rgb = RGBColor(43, 87, 151)
            run.font.size = Pt(28)

        # subtitle with date and author
        sub = doc.add_paragraph()
        sub_run = sub.add_run(f"Generated {datetime.now():%d %B %Y at %H:%M}")
        sub_run.font.size = Pt(10)
        sub_run.font.color.rgb = RGBColor(120, 120, 120)
        sub.add_run("\n")
        author_run = sub.add_run("Emmanuel Omego — Research Agent")
        author_run.font.size = Pt(10)
        author_run.font.color.rgb = RGBColor(120, 120, 120)

        # --- query section ---
        doc.add_heading("Research Query", level=1)
        q_para = doc.add_paragraph(query)
        q_para.style.font.size = Pt(11)

        # --- confidence ---
        conf = f.get("confidence", "low").upper()
        conf_para = doc.add_paragraph()
        conf_run = conf_para.add_run(f"Confidence Level: {conf}")
        conf_run.font.bold = True
        conf_run.font.size = Pt(11)
        if conf == "HIGH":
            conf_run.font.color.rgb = RGBColor(34, 139, 34)
        elif conf == "MEDIUM":
            conf_run.font.color.rgb = RGBColor(200, 150, 0)
        else:
            conf_run.font.color.rgb = RGBColor(180, 50, 50)

        # --- summary ---
        doc.add_heading("Summary", level=1)
        # use enhanced summary if we have it, otherwise the original
        summary_text = f.get("enhanced_summary", f.get("summary", "No summary available."))
        # split into paragraphs for readability
        for para_text in summary_text.split("\n\n"):
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                p.style.font.size = Pt(11)

        # --- data points table ---
        dps = f.get("data_points", [])
        if dps:
            doc.add_heading(f"Data Points ({len(dps)})", level=1)

            table = doc.add_table(rows=1, cols=4)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.style = "Light Grid Accent 1"

            # header row
            headers = ["Fact", "Value", "Unit", "Source"]
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = h
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.size = Pt(9)

            # data rows
            for dp in dps:
                row = table.add_row()
                row.cells[0].text = str(dp.get("fact", ""))
                row.cells[1].text = str(dp.get("value", ""))
                row.cells[2].text = str(dp.get("unit", ""))
                src = str(dp.get("source", ""))
                row.cells[3].text = src[:60] + "..." if len(src) > 60 else src
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(9)

        # --- sources ---
        sources = f.get("sources", [])
        if sources:
            doc.add_heading(f"Sources ({len(sources)})", level=1)
            for s in sources:
                p = doc.add_paragraph(s, style="List Bullet")
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(43, 87, 151)

        # --- gaps ---
        gaps = f.get("gaps", [])
        if gaps:
            doc.add_heading("Information Gaps", level=1)
            for g in gaps:
                p = doc.add_paragraph(g, style="List Bullet")
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(180, 50, 50)

        # --- metadata footer ---
        doc.add_paragraph()  # spacer
        footer_para = doc.add_paragraph()
        footer_run = footer_para.add_run("─" * 60)
        footer_run.font.color.rgb = RGBColor(200, 200, 200)
        footer_run.font.size = Pt(8)

        meta_para = doc.add_paragraph()
        mode = result.get("mode", "single")
        if mode == "parallel":
            meta_text = (f"Parallel mode: {m.get('num_workers','?')} workers, "
                        f"{m.get('total_turns','?')} total turns, {m.get('elapsed_seconds','?')}s")
        else:
            meta_text = (f"Single mode: {m.get('turns_used','?')} turns, "
                        f"{m.get('elapsed_seconds','?')}s")
        meta_text += f"\nModels: {m.get('planning_model','')} + {m.get('browser_model','')}"
        meta_text += f"\nTimestamp: {m.get('timestamp','')}"
        meta_run = meta_para.add_run(meta_text)
        meta_run.font.size = Pt(8)
        meta_run.font.color.rgb = RGBColor(150, 150, 150)

        doc.save(str(filepath))
        print(f"[DocumentationAgent] Word doc saved: {filepath}")
        return filepath

    # ─── PDF REPORT ──────────────────────────────────────────────────
    # same data but as a PDF — for people who just want to read it

    def _make_pdf(self, result: dict, filepath: Path) -> Path:
        """build a styled PDF from the research results."""
        f = result["findings"]
        m = result.get("metadata", {})
        query = result.get("query", "Unknown")

        # fpdf2's built-in Helvetica only supports latin-1 and has no
        # glyphs for C0/C1 control characters.  Strip those first, then
        # replace any remaining non-latin-1 chars with '?'.
        def safe(text: str) -> str:
            import re
            text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
            return text.encode("latin-1", errors="replace").decode("latin-1")

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.add_page()

        # title
        pdf.set_font("Helvetica", "B", 24)
        pdf.set_text_color(43, 87, 151)
        pdf.cell(0, 14, "Research Report", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 11)
        pdf.set_text_color(100, 100, 100)
        pdf.cell(0, 7, "Generated {0}".format(datetime.now().strftime("%d %B %Y at %H:%M")),
                 new_x="LMARGIN", new_y="NEXT")
        pdf.cell(0, 7, "Emmanuel Omego - Research Agent", new_x="LMARGIN", new_y="NEXT")

        # blue line
        pdf.ln(4)
        pdf.set_draw_color(43, 87, 151)
        pdf.set_line_width(0.8)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(8)

        # query
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_text_color(30, 30, 30)
        pdf.cell(0, 8, "Research Query", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 11)
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 6, safe(query))
        pdf.ln(4)

        # confidence
        conf = f.get("confidence", "low")
        colors = {"high": (34, 139, 34), "medium": (200, 150, 0), "low": (180, 50, 50)}
        r, g, b = colors.get(conf, (100, 100, 100))
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(r, g, b)
        pdf.cell(0, 7, "Confidence: {0}".format(conf.upper()), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

        # summary (use enhanced if available)
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_text_color(30, 30, 30)
        pdf.cell(0, 8, "Summary", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 11)
        pdf.set_text_color(50, 50, 50)
        summary = f.get("enhanced_summary", f.get("summary", "No summary."))
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 6, safe(summary))
        pdf.ln(6)

        # data table
        dps = f.get("data_points", [])
        if dps:
            pdf.set_font("Helvetica", "B", 13)
            pdf.set_text_color(30, 30, 30)
            pdf.cell(0, 8, "Data Points ({0})".format(len(dps)), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)

            # header
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_fill_color(43, 87, 151)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(65, 8, "Fact", border=1, fill=True)
            pdf.cell(40, 8, "Value", border=1, fill=True)
            pdf.cell(20, 8, "Unit", border=1, fill=True)
            pdf.cell(65, 8, "Source", border=1, fill=True, new_x="LMARGIN", new_y="NEXT")

            # rows
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(50, 50, 50)
            for i, dp in enumerate(dps):
                stripe = i % 2 == 0
                if stripe:
                    pdf.set_fill_color(240, 245, 250)
                fact = safe(str(dp.get("fact", "")))[:45]
                val = safe(str(dp.get("value", "")))[:28]
                unit = safe(str(dp.get("unit", "")))[:12]
                src = safe(str(dp.get("source", "")))
                if len(src) > 45:
                    src = src[:42] + "..."
                pdf.cell(65, 7, fact, border=1, fill=stripe)
                pdf.cell(40, 7, val, border=1, fill=stripe)
                pdf.cell(20, 7, unit, border=1, fill=stripe)
                pdf.cell(65, 7, src, border=1, fill=stripe, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(6)

        # sources
        sources = f.get("sources", [])
        if sources:
            pdf.set_font("Helvetica", "B", 13)
            pdf.set_text_color(30, 30, 30)
            pdf.cell(0, 8, "Sources ({0})".format(len(sources)), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(43, 87, 151)
            for s in sources:
                display = s if len(s) < 120 else s[:117] + "..."
                pdf.set_x(pdf.l_margin)
                pdf.multi_cell(0, 5, safe(display))
            pdf.ln(4)

        # gaps
        gaps = f.get("gaps", [])
        if gaps:
            pdf.set_font("Helvetica", "B", 13)
            pdf.set_text_color(180, 50, 50)
            pdf.cell(0, 8, "Gaps", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(80, 80, 80)
            for g in gaps:
                gt = g if len(g) < 150 else g[:147] + "..."
                pdf.set_x(pdf.l_margin)
                pdf.multi_cell(0, 6, safe("  - {0}".format(gt)))
            pdf.ln(4)

        # footer
        pdf.ln(6)
        pdf.set_draw_color(200, 200, 200)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(4)
        pdf.set_font("Helvetica", "I", 9)
        pdf.set_text_color(130, 130, 130)
        mode = result.get("mode", "single")
        if mode == "parallel":
            pdf.cell(0, 5, "Parallel: {0} workers, {1} turns, {2}s".format(
                m.get('num_workers', '?'), m.get('total_turns', '?'), m.get('elapsed_seconds', '?')),
                new_x="LMARGIN", new_y="NEXT")
        else:
            pdf.cell(0, 5, "Single: {0} turns, {1}s".format(
                m.get('turns_used', '?'), m.get('elapsed_seconds', '?')),
                new_x="LMARGIN", new_y="NEXT")
        pdf.cell(0, 5, "Models: {0} + {1}".format(
            m.get('planning_model', ''), m.get('browser_model', '')),
            new_x="LMARGIN", new_y="NEXT")

        pdf.output(str(filepath))
        print(f"[DocumentationAgent] PDF saved: {filepath}")
        return filepath
